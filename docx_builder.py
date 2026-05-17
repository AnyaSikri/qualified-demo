"""Builders for Word document + sources JSON outputs.

Both take the structured `narrative_data` dict produced by
generator.generate_narrative(). The docx mirrors the narrative the
reviewer reads; the JSON mirrors the citation trail the auditor reads.
"""

import json

from docx import Document
from docx.shared import Pt


def _fmt_adae_cite(adae_source):
    cols = ", ".join(adae_source["columns_used"]) or "(no fields)"
    return f"ADAE row {adae_source['row']} ({cols})"


def _fmt_protocol_cite(protocol_source):
    pages = ",".join(str(p) for p in protocol_source["pages"]) or "?"
    return f"Protocol § {protocol_source['section_title']} p.{pages}"


def _format_source_line(sources):
    parts = [_fmt_adae_cite(a) for a in sources["adae"]]
    parts += [_fmt_protocol_cite(p) for p in sources["protocol"]]
    return "; ".join(parts) if parts else "(no sources)"


def build_narrative_docx(narrative_data, out_path):
    """Write a Word document of the narrative with per-section source lines."""
    doc = Document()

    # Body font defaults — 11pt Calibri is python-docx default; bump slightly.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(
        f"Patient Safety Narrative — Subject {narrative_data['subject_id']}",
        level=0,
    )
    title.alignment = 1  # center

    sub = doc.add_paragraph()
    sub.add_run(
        f"Template: {narrative_data['template']['name']} "
        f"v{narrative_data['template']['version']}"
    ).italic = True

    for section in narrative_data["sections"]:
        doc.add_heading(section["title"], level=1)
        doc.add_paragraph(section["text"])

        sources_para = doc.add_paragraph()
        label = sources_para.add_run("Sources: ")
        label.bold = True
        sources_para.add_run(_format_source_line(section["sources"])).italic = True

    doc.save(out_path)
    return out_path


def build_sources_json(narrative_data, out_path):
    """Write a JSON file capturing the full citation trail for the narrative."""
    payload = {
        "subject_id": narrative_data["subject_id"],
        "template": narrative_data["template"],
        "sections": [
            {
                "section_id": s["section_id"],
                "title": s["title"],
                "text": s["text"],
                "sources": s["sources"],
            }
            for s in narrative_data["sections"]
        ],
    }
    with open(out_path, "w") as f:
        json.dump(payload, f, indent=2, default=str)
    return out_path


# ---------------------------------------------------------------------------
# Standalone smoke test — builds both outputs from the JSON written by
# `python generator.py SUBJ-1042` so this can run without re-billing GPT-4o.
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    import os
    import sys

    here = os.path.dirname(os.path.abspath(__file__))
    subject_id = sys.argv[1] if len(sys.argv) > 1 else "SUBJ-1042"
    json_in = os.path.join(here, "data", f"narrative_{subject_id}.json")

    if not os.path.exists(json_in):
        sys.exit(
            f"{json_in} not found. Run `python generator.py {subject_id}` first."
        )

    with open(json_in) as f:
        narrative = json.load(f)

    docx_out = os.path.join(here, "data", f"narrative_{subject_id}.docx")
    sources_out = os.path.join(here, "data", f"sources_{subject_id}.json")

    build_narrative_docx(narrative, docx_out)
    build_sources_json(narrative, sources_out)

    print(f"wrote {docx_out} ({os.path.getsize(docx_out)} bytes)")
    print(f"wrote {sources_out} ({os.path.getsize(sources_out)} bytes)")
